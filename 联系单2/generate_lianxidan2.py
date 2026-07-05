# -*- coding: utf-8 -*-
"""
生成《工作联系单2（各园合并版）.docx》
依据：
  1) 工程量确认单合并版（最终修订版，2026.7.4）
  2) 已盖章的变更联系单1（扫描件，6所幼儿园）
  3) 工作联系单范本
规则：确认单条目剔除联系单1已确认内容后，剩余条目按范本格式生成各校联系单2，
      合并在同一份文件中。
运行：pip install python-docx && python3 generate_lianxidan2.py
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_NAME = "东升、东红、千秋、东平、南俸、上埇和机关等7所幼儿园室外配套附属设施建设工程"

SCHOOLS = [
    {
        "name": "南俸幼儿园",
        "user_unit": "琼海市石壁镇中心幼儿园",
        "items": [
            "原设计南侧大门外沥青路配套排水沟不再实施；部分路缘石不再实施，长度180m；下坡处防护栏长度明确为35m（沥青路不再实施及对现有损坏区域进行处理已按联系单1执行）。",
            "原联系单1中EPDM地板调整为“右侧EPDM、中间硬化、左侧悬浮地板”三个区域；现进一步调整为：取消EPDM区域，整体调整为硬化区域1697.85㎡及悬浮式地板区域1599㎡，悬浮地板规格为30cm×30cm×1.6cm。",
            "沿红线点1-2-3-4-5-6-7-8新增栏杆围墙232m（原联系单1表述为增加铁艺围墙，现明确围墙形式及长度）；取消LOGO墙，方量4.5m³。",
            "变压器周边增设成品隔离护栏18m；原联系单1中提到的高压计量设备取消（变压器按250KVA施工及增加电线杆已按联系单1执行，电线杆明确为8m高砼电线杆1根）。",
            "围墙边四周绿化位置增加成品木栅栏，长度168.8m。",
            "根据南俸幼儿园监控系统使用需求，增加以下配套设备：16路双盘位录像机1台、10T硬盘1块、8口POE交换机2台、8口千兆交换机1台、显示器1台。",
        ],
    },
    {
        "name": "东红幼儿园",
        "user_unit": "琼海市大路镇中心幼儿园东红分园",
        "items": [
            "保安亭布局及使用功能同步优化后，原配套人行通道不锈钢门及屋顶光伏系统不再另行实施（保安亭宽敞布局调整已按联系单1执行，做法参照东升幼儿园保安亭）。",
            "原联系单1中曲面山包变更为成品攀爬玩具；现成品攀爬玩具取消，曲面山包取消后该区域面层调整为悬浮式地板（面积已计入悬浮地板833.16㎡合计中）。",
            "悬浮式地板铺设面积明确为833.16㎡，下部硬化工序为：拉毛、切缝宽6mm、深5cm、填充沥青砂浆、棉毡养护（悬浮地板规格变更为30cm×30cm×1.6cm已按联系单1执行）。",
            "根据实际施工要求，新购置POE交换机1台（海康DS-XS10-P，8口百兆PoE电口）；录像机沿用学校室内已有设备，不再另购；本项不涉及新增硬盘及显示器。",
            "教学楼门口增加铸铁盖板水沟22.5m，内径25cm宽、20cm高，两侧砌单砖抹灰。",
        ],
    },
    {
        "name": "东升幼儿园",
        "user_unit": "琼海市万泉镇中心幼儿园",
        "items": [
            "原设计PVC运动地板735.22㎡变更为减震悬浮地板已按联系单1执行，现明确铺设面积664㎡；硬化区域调整为200厚级配碎石+200厚C25混凝土，硬化面积725㎡，硬化优化工序为：拉毛、切缝宽6mm、深5cm、填充沥青砂浆、棉毡养护。",
            "隔油池（2T成品玻璃钢，已按联系单1执行）埋置16.79m DN200波纹管接至W8污水井；旁边增设两个铸铁篦子雨水口，通过5.4m DN200波纹管排至雨水井；原联系单1中提到的洗手池水龙头、大理石小径、花池排水井及大门绿化优化内容现场未实施，取消。",
            "原联系单1中建议增设的戏水池顶棚现场未实施，取消。",
            "应使用单位要求，拆除单层混合结构建筑；拆除清理后，局部场地做硬化并铺设悬浮式地板，面积70㎡，拼接图案由厂家提供；其余清理后未硬化区域作为绿化场地处理。",
            "取消原电动伸缩门1座，替换为成品双扇咖色氟碳漆热镀锌方管大门，高2.5m、总宽6m；大门外指定场地及园内拆除清理后未硬化区域，换填300厚种植土并铺设台湾草，面积438㎡。",
            "大门口消防车道面积103.79㎡，原已新建完成，后因现场使用及通行调整需要，对已完成消防车道进行拆除后重新施工；大门内消防车道坡道两侧增加挡土墙，总长22.97m；坡道两侧新增栏杆扶手，总长度15.5m。",
            "应业主要求，对新建围墙做法进行调整，其中总高2.9m的围墙长105.05m，总高2.5m的围墙长29.94m。",
            "原单层混合结构建筑门口位置新增铸铁盖板14.4m。",
            "取消所有教学楼监控。",
        ],
    },
    {
        "name": "千秋幼儿园",
        "user_unit": "琼海市塔洋镇中心幼儿园",
        "items": [
            "活动场地除悬浮地板外的其余铺装场地均简化为硬化，硬化区域调整为200厚级配碎石+200厚C25混凝土，硬化优化工序为：拉毛、切缝宽6mm、深5cm、填充沥青砂浆、棉毡养护（悬浮地板规格变更为30cm×30cm×1.6cm、铺设总面积607.2㎡已按联系单1执行）。",
            "原联系单1中建议增加的学校全称广告牌现场未实施，取消。",
            "按业主意见，取消南向内部园门及底部台阶；取消两侧3m高围挡设施48㎡。",
            "根据现场已施工实际情况，取消成品木花箱4个。",
            "围墙造型柱及高度调整：取消原围墙500×500造型柱；将原2m高围墙分段调整为不同高度施工，其中2.65m高围墙18m、3.3m高围墙20m、3.05m高围墙35.22m；均为18墙双面抹灰，内外双侧一底两面真石漆涂料。",
            "西侧楼梯后新增2.5m高围墙，总长6m，18墙双面抹灰，内外双侧一底两面真石漆涂料。",
            "新增沙池水池顶部遮雨棚，面积53.68㎡，做法按变更图为准。",
            "原趣味花架材料规格进行变更，均采用热镀锌钢管，面喷咖色氟碳漆，做法按变更图为准。",
            "根据实际施工要求，新购置录像机1台（海康DS-7808N-Q1）、硬盘1块（希捷6T）、配套显示器1台（佩奇PQ-L324KA）、POE交换机1台（海康DS-XS10-P，8口百兆PoE电口）。",
        ],
    },
    {
        "name": "机关幼儿园",
        "user_unit": "琼海市机关幼儿园",
        "items": [
            "变压器由联系单1中确认的315KVA调整为按原设计图纸250KVA施工；原联系单1中提到的高压计量设备取消（供电端增加8m电线杆、变压器周边隔离护栏18m及草皮绿化已按联系单1执行）。",
            "根据实际施工要求，新购置录像机1台（海康DS-7808N-R2，两盘位）、硬盘1块（希捷10T）、配套显示器1台（佩奇PQ-L324KA）、POE交换机2台（海康DS-XS10-P，8口百兆PoE电口，因场地特殊需分2段布线）。",
        ],
    },
    {
        "name": "上埇幼儿园",
        "user_unit": "琼海市嘉积镇上埇幼儿园",
        "items": [
            "根据使用单位要求，原设计大门位置不再调整，仍保留在原位置使用；原联系单1中提出的大门位置调整及铁艺大门相关内容现场未实施，取消。",
            "已建值班室至大门右边5段围墙做法由联系单1中的铝合金围墙（43.2m）调整为铁艺围墙，长度68.4m；其中18m曾建成砖体基础、砌砖、刷涂料，因使用单位原因铁艺栏杆未安装，后该18m已拆除。",
            "原大门口黄锈石铺装取消，调整为现浇硬化，硬化面积951㎡；硬化做法为200厚级配碎石+200厚C25混凝土，硬化工序为：拉毛、切缝宽6mm、深5cm、填充沥青砂浆、棉毡养护；原联系单1中提到的草皮绿化现场未实施，取消。",
            "新建筑背面新增挡土墙18m，墙顶相对标高0.46～1.8m。",
            "大门位置及场地西南侧小门位置各新增电动伸缩门1樘，共计2樘电动伸缩门。",
            "新增成品保安亭1座，规格为长4m、宽3m、高3.3m，并做好防雷接地。",
            "新建片石挡土墙上面的围墙，长度65m，该段围墙无砖基础（直接建于片石挡土墙顶部）。",
            "由于现场将化粪池埋置在场地西北角，DN200污水管增加32.58m，污水检查井增加1个。",
            "由于市政接入点改变，原主电缆YJV22-4×35/SC50/FC线路总长度由28.65m调整为114.45m，并增设手孔井1座，做法不变。",
            "原设计场地中央沙池调整为戏水池，尺寸不变；同时取消沙水区遮阳棚75㎡。",
            "大门坡底截水沟总长度调整为28m；东侧挡土墙旁新增排水沟37m，采用砼制水沟盖板，规格为500×400×55厚；东侧挡土墙旁新增100厚C25混凝土硬化17.48㎡。",
            "西南侧小门外增设DN400雨水管，长度9.4m；雨水管端部设置800×800、深1m沉沙井1座，井壁采用24墙砌筑抹灰；沉沙井安装成品铸铁井盖。",
            "拆除大门前停车场损坏彩砖；新建彩砖规格为400×400×50厚，总面积47㎡。",
            "根据实际施工要求，新购置POE交换机2台（海康DS-XS10-P，8口百兆PoE电口）；录像机沿用学校室内已有设备，不再另购；本项不涉及新增硬盘及显示器。",
            "取消场地中央挡土墙外立面火山岩装饰层133.45㎡。",
            "场地中间挡土墙顶护栏做法调整：原设计玻璃栏杆护栏236.85m，调整为不锈钢栏杆扶手护栏220m；做法参照东升学校场地中间坡道挡土墙两侧不锈钢护栏做法。",
            "因使用单位原因，大门口斜坡位置新建的草皮完成后又予以铲除，面积150㎡。",
        ],
    },
    {
        "name": "东平幼儿园",
        "user_unit": "琼海市东平幼儿园",
        "items": [
            "根据东平幼儿园监控系统使用需求，增加以下配套设备：16路双盘位录像机1台、10T硬盘1块、POE交换机2台、8口千兆交换机1台、显示器/监视器1台。",
        ],
    },
]

OPINION_BLOCKS = [
    ("监理单位意见：", ["监理单位（盖章）：", "总监理工程师（签字）：", "日期："]),
    ("使用单位意见：", ["使用单位（盖章）：", "项目负责人（签字）：", "日期："]),
    ("设计单位意见：", ["设计单位（盖章）：", "项目负责人（签字）：", "日期："]),
    ("建设单位意见：", ["建设单位（盖章）：", "项目负责人（签字）：", "日期："]),
    ("教育局意见：", ["琼海市教育局（盖章）：", "项目负责人（签字）：", "日期："]),
]


def set_font(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_para(cell_or_doc, text, size=10.5, bold=False, align=None,
             space_after=2, first_line_indent=None):
    if hasattr(cell_or_doc, "add_paragraph"):
        p = cell_or_doc.add_paragraph()
    else:
        p = cell_or_doc
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
    return p


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tcPr.append(borders)


def clear_cell(cell):
    cell.paragraphs[0].text = ""
    return cell


def build_school_sheet(doc, school, first=False):
    if not first:
        doc.add_page_break()

    add_para(doc.add_paragraph(), "工作联系单", size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc.add_paragraph(),
             f"工程名称：{PROJECT_NAME}　　编号：",
             size=10.5, space_after=4)

    table = doc.add_table(rows=2 + len(OPINION_BLOCKS), cols=2)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(11.0)
        for cell in row.cells:
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # 第1行：致（合并两列）
    head_cell = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    clear_cell(head_cell)
    add_para(head_cell, "致：海口市工程监理有限公司　　　　　　　　（监理单位）")
    add_para(head_cell, f"　　{school['user_unit']}　　　　（使用单位）")
    add_para(head_cell, "　　中昌设计集团有限公司　　　　　　　　　（设计单位）")
    add_para(head_cell, "　　琼海市城市投资运营有限公司　　　　　　（建设单位）")
    add_para(head_cell, "　　琼海市教育局")

    # 第2行：事由 + 承包单位（合并两列）
    body_cell = table.rows[1].cells[0].merge(table.rows[1].cells[1])
    clear_cell(body_cell)
    add_para(body_cell, "事由：", bold=True)
    add_para(body_cell,
             "根据现场实际情况及使用单位提出的若干意见，在已确认的工作联系单1基础上，"
             "对原设计图纸相关内容做如下变更（联系单1已确认的内容不再重复列入）：",
             first_line_indent=21)
    for i, item in enumerate(school["items"], start=1):
        add_para(body_cell, f"{i}.{item}", first_line_indent=21)
    add_para(body_cell, "请各参建单位予以确认。", first_line_indent=21)
    add_para(body_cell, "")
    add_para(body_cell, "　　　　　　　　　　　　　　　　　　　　承包单位：")
    add_para(body_cell, "　　　　　　　　　　　　　　　　　　　　项目经理：")
    add_para(body_cell, "　　　　　　　　　　　　　　　　　　　　日　期　：")

    # 各单位意见行
    for idx, (label, right_lines) in enumerate(OPINION_BLOCKS):
        row = table.rows[2 + idx]
        left = clear_cell(row.cells[0])
        add_para(left, label)
        for _ in range(3):
            add_para(left, "")
        right = clear_cell(row.cells[1])
        add_para(right, "")
        for line in right_lines:
            add_para(right, line, space_after=8)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    for i, school in enumerate(SCHOOLS):
        build_school_sheet(doc, school, first=(i == 0))

    out = "工作联系单2（各园合并版）.docx"
    doc.save(out)
    print(f"已生成：{out}")


if __name__ == "__main__":
    main()
