# -*- coding: utf-8 -*-
"""
生成：机关、上埇、东红、南俸、千秋 5所幼儿园的
1. 工作联系单（二）—— 补充联系单一未覆盖、但确认单存在的内容
2. 工程量确认单（新）—— 按联系单顺序重排、差异内容加注

版式与《东升幼儿园》两份文件（gen_docs.py）保持一致，1:1复刻原件模板。
东平幼儿园本轮不处理。
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "仿宋_GB2312"
BODY_SIZE = 12
TITLE_SIZE = 24
CONTENT_WIDTH_CM = 17.33

PROJECT_FULL_NAME = "东升、东红、千秋、东平、南俸、上埇和机关等7所幼儿园室外配套附属设施建设工程"


def set_east_asia_font(run, size=BODY_SIZE, bold=False, name=FONT_NAME):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def set_single_spacing(paragraph, space_before=0, space_after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{tag}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        mar.append(node)
    tcPr.append(mar)


def new_section_margins(doc, top=1.6, bottom=1.6, left=1.84, right=1.84):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_title(doc, text, space_after_pt=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p, space_before=0, space_after=space_after_pt)
    run = p.add_run(text)
    set_east_asia_font(run, size=TITLE_SIZE, bold=False)
    return p


def add_plain_para(doc, text, size=BODY_SIZE, indent_chars=0, align=None,
                    space_before=0, space_after=0):
    p = doc.add_paragraph()
    set_single_spacing(p, space_before, space_after)
    if indent_chars:
        p.paragraph_format.first_line_indent = Pt(size * indent_chars)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_east_asia_font(run, size=size)
    return p


def set_cell(cell, text, size=BODY_SIZE, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             valign=WD_ALIGN_VERTICAL.CENTER, multi_line=None):
    cell.text = ""
    set_cell_margins(cell)
    cell.vertical_alignment = valign
    lines = multi_line if multi_line else [text]
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        set_single_spacing(p, 0, 0)
        p.alignment = align
        run = p.add_run(line)
        set_east_asia_font(run, size=size, bold=bold)
    return cell


def set_row_height(row, cm_value, rule=WD_ROW_HEIGHT_RULE.AT_LEAST):
    row.height = Cm(cm_value)
    row.height_rule = rule


# ------------------------------------------------------------------
# 通用：工作联系单（二）
# ------------------------------------------------------------------

def build_lianxidan_2(*, school, use_unit, lianxi_bianhao, ref_bianhao,
                       intro_text, items, output_path):
    doc = Document()
    new_section_margins(doc)

    add_title(doc, "工作联系单", space_after_pt=14)

    p = doc.add_paragraph()
    set_single_spacing(p, 0, 10)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(CONTENT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(f"工程名称：{PROJECT_FULL_NAME}\t编号：{lianxi_bianhao}")
    set_east_asia_font(run, size=BODY_SIZE)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col0_w, col1_w = Cm(6.0), Cm(CONTENT_WIDTH_CM - 6.0)
    for row in table.rows:
        row.cells[0].width = col0_w
        row.cells[1].width = col1_w

    top_cell = table.cell(0, 0).merge(table.cell(0, 1))
    top_cell.text = ""
    set_cell_margins(top_cell, top=120, bottom=120, left=140, right=140)
    top_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    def top_para(first=False):
        p = top_cell.paragraphs[0] if first else top_cell.add_paragraph()
        set_single_spacing(p, 0, 0)
        return p

    p = top_para(first=True)
    pf = p.paragraph_format
    pf.left_indent = Cm(1.1)
    pf.first_line_indent = Cm(-1.1)
    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Cm(7.2))
    units = [
        ("致：海口市工程监理有限公司", "（监理单位）"),
        (use_unit, "（使用单位）"),
        ("中昌设计集团有限公司", "（设计单位）"),
        ("琼海市城市投资运营有限公司", "（建设单位）"),
        ("琼海市教育局", ""),
    ]
    for i, (name, label) in enumerate(units):
        run = p.add_run(f"{name}\t{label}")
        set_east_asia_font(run, size=BODY_SIZE)
        if i != len(units) - 1:
            run.add_break()

    p = top_para()
    set_single_spacing(p, 8, 0)
    run = p.add_run("事由：")
    set_east_asia_font(run, size=BODY_SIZE)

    p = top_para()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(BODY_SIZE * 2)
    run = p.add_run(intro_text)
    set_east_asia_font(run, size=BODY_SIZE)

    for i, text in enumerate(items, start=1):
        p = top_para()
        run = p.add_run(f"{i}.{text}")
        set_east_asia_font(run, size=BODY_SIZE)

    p = top_para()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(BODY_SIZE * 2)
    set_single_spacing(p, 6, 0)
    run = p.add_run("请各参建单位予以确认。")
    set_east_asia_font(run, size=BODY_SIZE)

    for label in ("承包单位：", "项目经理：", "日\u3000\u3000期："):
        p = top_para()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_single_spacing(p, 4, 0)
        pf = p.paragraph_format
        pf.right_indent = Cm(2.0)
        run = p.add_run(label)
        set_east_asia_font(run, size=BODY_SIZE)

    opinion_rows = [
        ("监理单位意见：", ["监理单位（盖章）：", "总监理工程师（签字）：", "日期："]),
        ("使用单位意见：", ["使用单位（盖章）：", "项目负责人（签字）：", "日期："]),
        ("设计单位意见：", ["设计单位（盖章）：", "项目负责人（签字）：", "日期："]),
        ("建设单位意见：", ["建设单位（盖章）：", "项目负责人（签字）：", "日期："]),
        ("教育局意见：", ["琼海市教育局（盖章）：", "项目负责人（签字）：", "日期："]),
    ]
    for i, (left_label, right_lines) in enumerate(opinion_rows, start=1):
        row = table.rows[i]
        set_row_height(row, 2.7)
        set_cell(row.cells[0], "", valign=WD_ALIGN_VERTICAL.TOP, multi_line=[left_label])
        set_cell(row.cells[1], "", valign=WD_ALIGN_VERTICAL.TOP, multi_line=right_lines)

    doc.save(output_path)
    print(f"{output_path} done")


# ------------------------------------------------------------------
# 通用：工程量确认单（新）
# ------------------------------------------------------------------

def build_quedingdan_new(*, school, project_name_full, content_name, position,
                          old_bianhao, note_text, shiyou_text, items, output_path):
    doc = Document()
    new_section_margins(doc)

    add_title(doc, "工程量确认单", space_after_pt=16)

    p = doc.add_paragraph()
    set_single_spacing(p, 0, 10)
    run = p.add_run(note_text)
    set_east_asia_font(run, size=10.5)

    info = doc.add_table(rows=2, cols=4)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    widths = [Cm(2.88), Cm(7.31), Cm(1.67), Cm(5.47)]
    for row in info.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    set_row_height(info.rows[0], 3.0)
    set_row_height(info.rows[1], 1.69)

    set_cell(info.rows[0].cells[0], "工程名称", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(info.rows[0].cells[1], project_name_full, align=WD_ALIGN_PARAGRAPH.LEFT,
              valign=WD_ALIGN_VERTICAL.TOP)
    set_cell(info.rows[0].cells[2], "内容", align=WD_ALIGN_PARAGRAPH.CENTER,
              valign=WD_ALIGN_VERTICAL.TOP)
    set_cell(info.rows[0].cells[3], content_name, align=WD_ALIGN_PARAGRAPH.LEFT,
              valign=WD_ALIGN_VERTICAL.TOP)

    set_cell(info.rows[1].cells[0], "位置", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(info.rows[1].cells[1], position, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(info.rows[1].cells[2], "", align=WD_ALIGN_PARAGRAPH.CENTER,
              multi_line=["计量", "编号"])
    set_cell(info.rows[1].cells[3], f"{old_bianhao}（修订）", align=WD_ALIGN_PARAGRAPH.LEFT)

    add_plain_para(doc, f"事由：{shiyou_text}", space_before=6, space_after=0)
    add_plain_para(doc, "根据会议纪要，结合使用单位提出的若干意见做出修改，修改内容如下：",
                   indent_chars=2, space_before=0, space_after=4)

    for i, text in enumerate(items, start=1):
        add_plain_para(doc, f"{i}.{text}", space_before=0, space_after=4)

    add_plain_para(doc, "以上变更工程量详见后附计算表。", space_before=6, space_after=8)

    bottom = doc.add_table(rows=1, cols=4)
    bottom.style = "Table Grid"
    bottom.alignment = WD_TABLE_ALIGNMENT.CENTER
    bottom.autofit = False
    widths2 = [Cm(4.19), Cm(4.20), Cm(4.47), Cm(4.48)]
    labels = ["建设单位意见：", "代管单位意见：", "监理单位意见：", "施工单位意见："]
    for cell, w in zip(bottom.rows[0].cells, widths2):
        cell.width = w
    set_row_height(bottom.rows[0], 4.32)
    for cell, label in zip(bottom.rows[0].cells, labels):
        set_cell(cell, label, valign=WD_ALIGN_VERTICAL.TOP)

    doc.save(output_path)
    print(f"{output_path} done")


OUT_DIR = "/workspace/output_docs"

# ====================================================================
# 1. 机关幼儿园
# ====================================================================
build_lianxidan_2(
    school="机关幼儿园",
    use_unit="琼海市机关幼儿园",
    lianxi_bianhao="2",
    ref_bianhao="1",
    intro_text=(
        "根据现场实际情况及使用单位提出的若干意见，对原设计图纸做出部分修改，"
        "以下内容为对工作联系单（编号1）的补充说明，工作联系单（编号1）与本联系单"
        "（编号2）合并后与本项目机关幼儿园工程量确认单内容一一对应："
    ),
    items=[
        "由于幼儿园内已有路灯设施，故决定取消新建路灯8台，及电缆设施343.92m。",
        "根据使用单位要求新购置录像机、10T机械硬盘、配套显示器各1台，POE交换机2台。",
    ],
    output_path=f"{OUT_DIR}/机关幼儿园_工作联系单（二）.docx",
)

build_quedingdan_new(
    school="机关幼儿园",
    project_name_full=f"{PROJECT_FULL_NAME}-机关幼儿园室外配套附属设施建设工程",
    content_name="机关幼儿园增加部分改造工程",
    position="机关幼儿园",
    old_bianhao="005",
    note_text=(
        "说明：本确认单根据《工作联系单》（编号1，已签字盖章）与《工作联系单》"
        "（编号2，补充）整理调整，替代原《工程量确认单》计量编号005（机关幼儿园）。"
        "条目顺序已按工作联系单顺序重排，与工作联系单内容不一致或联系单中已取消的内容，"
        "均在对应条目下加注说明。"
    ),
    shiyou_text="关于机关幼儿园增加部分改造工程的确认事宜",
    items=[
        "原设计图悬浮地板规格25cm×25cm×1.3cm，590㎡，变更为规格为30cm×30cm×1.6cm，"
        "590㎡，硬化区域现调整硬化为200厚级配碎石+200厚C25混凝土，硬化优化工序：拉毛."
        "切缝宽6mm，深5cm.填充沥青砂浆.棉毡养护；（对应工作联系单一第1条）",

        "铲除外运原地面遗留的塑胶层212.8平方米；（对应工作联系单一第2条）",

        "因原位置不够，新建变压器位置改至场地南北侧角落原种植区，变压器安装成品隔离护"
        "栏18m，四周种植草皮(最终种植面积根据施工资料确定）；供电端增加一根8m电线杆；"
        "（注：工作联系单一第3条提到的变压器容量315KVA及高压计量设备为设备参数，"
        "不纳入本次工程量确认范围，具体以设备供应商及供电局要求为准。）",

        "原设计中沙池的回填沙量新增加6立方米；（对应工作联系单一第4条）",

        "由于幼儿园内已有路灯设施，故决定取消新建路灯8台，及电缆设施343.92m；"
        "（对应工作联系单二第1条）",

        "根据使用单位要求新购置录像机、10T机械硬盘、配套显示器各1台，POE交换机2台。"
        "（对应工作联系单二第2条）",
    ],
    output_path=f"{OUT_DIR}/机关幼儿园_工程量确认单（新）.docx",
)

# ====================================================================
# 2. 上埇幼儿园
# ====================================================================
build_lianxidan_2(
    school="上埇幼儿园",
    use_unit="琼海市嘉积镇上埇幼儿园",
    lianxi_bianhao="2",
    ref_bianhao="1",
    intro_text=(
        "根据现场实际情况及使用单位提出的若干意见，对原设计图纸做出部分修改，"
        "以下内容为对工作联系单（编号1）的补充说明，工作联系单（编号1）与本联系单"
        "（编号2）合并后与本项目上埇幼儿园工程量确认单内容一一对应："
    ),
    items=[
        "由于现场将化粪池埋置在场地西北角，致DN200污水管总长度变更为108.67m，"
        "污水检查井总数变更为8个。",
        "原大门口整片黄锈石铺装取消1633.9㎡，下部仅做现浇硬化面积951㎡，硬化区域现调整"
        "硬化为200厚级配碎石+200厚C25混凝土（拉毛.切缝宽6mm，深5cm.填充沥青砂浆.棉毡"
        "养护）。",
        "由于市政接入点改变，原主电缆YJV22-4X35/SC50/FC线路总长度28.65m变更为114.45m，"
        "并增设一个手孔井做法不变。",
        "场地中央原沙池变更为戏水池，尺寸不变。",
        "大门坡底截水沟总长度变更为28m；东侧挡土墙底部新增排水沟37m，砼制水沟盖板"
        "500*400*55厚；东侧挡土墙底新增100厚C25混凝土硬化17.48㎡。",
        "拆除大门前停车场损坏的彩砖，并新建彩砖：规格为400*400*50厚，总面积47㎡。",
        "根据使用单位要求新购置录像机、POE交换机各1台。",
        "根据现场实际施工情况取消场地中央挡土墙二外立面的火山岩装饰层133.45㎡；"
        "取消沙水区的遮阳棚75㎡。",
        "根据现场实际施工情况，将场地中间挡土墙顶的玻璃栏杆护栏做法由236.85m变更为"
        "不锈钢栏杆扶手护栏220m。做法参照东升学校场地中间坡道挡土墙两侧不锈钢护栏做法。",
    ],
    output_path=f"{OUT_DIR}/上埇幼儿园_工作联系单（二）.docx",
)

build_quedingdan_new(
    school="上埇幼儿园",
    project_name_full=f"{PROJECT_FULL_NAME}-嘉积镇上埇幼儿园室外配套附属设施建设工程",
    content_name="上埇幼儿园增加部分改造工程",
    position="上埇幼儿园",
    old_bianhao="006",
    note_text=(
        "说明：本确认单根据《工作联系单》（编号1，已签字盖章）与《工作联系单》"
        "（编号2，补充）整理调整，替代原《工程量确认单》计量编号006（上埇幼儿园）。"
        "条目顺序已按工作联系单顺序重排，与工作联系单内容不一致或联系单中已取消的内容，"
        "均在对应条目下加注说明；本单为该校本轮问题最复杂的一份，部分条目因联系单与"
        "确认单差异较大，已单独加注说明，请重点复核。"
    ),
    shiyou_text="关于上埇幼儿园增加部分改造工程的确认事宜",
    items=[
        "拆除原大门，新增成品保安亭1座(做好防雷接地），规格为长4m、宽3m、高3.3m；"
        "（注：工作联系单一第1条为“大门右边的5-6段围墙变更为铁艺大门”，与本条内容"
        "不一致，最终实施情况需现场核实确认，暂按本条执行。）",

        "根据现场施工情况，新建筑背面即场地南侧新增镀锌管围墙长44.30m，新大门东侧"
        "挡土墙顶新增镀锌管围墙18.78m；（注：工作联系单一第2条原为“大门对面1-2-3-4段"
        "围墙共40米，变更为铝合金围墙”，经设计单位优化调整，材质变更为镀锌管围墙，"
        "长度据实调整为63.08m。）",

        "场地东侧挡土墙顶新增镀锌管围墙100.59m(据实结算，仅在原挡土墙顶安装栏杆构件）；"
        "（对应工作联系单一第3条“已建值班室至大门右边的5段围墙”，经设计单位优化调整。）",

        "原活动场EPDM面层1101.01㎡变更为悬浮式地板，悬浮地板总铺设面积为1808㎡，下部"
        "硬化面积同此）；悬浮地板调整为30cm×30cm×1.6cm，590㎡，硬化区域现调整硬化为"
        "200厚级配碎石+200厚C25混凝土，硬化优化工序：拉毛.切缝宽6mm，深5cm.填充沥青"
        "砂浆.棉毡养护；（对应工作联系单一第4条）",

        "场地西南角增设成品12m宽电动折叠门，大门外增设9.4m长DN400雨水管，端部设置"
        "800*800深1m沉沙井，井壁24墙砌筑抹灰，安装成品铸铁井盖；根据学校要求，大门"
        "位置外移，该部位原已施工完成的新建围墙18.7m（高0.45m、18墙砌筑并抹灰）现需"
        "全部拆除，已完成工程量按上述数量计列；（对应工作联系单一第5条“拆除原有大门后"
        "新建铝艺围墙，原斜坡车行区域改造为草皮绿化，新大门设于原大门右侧”，其中围墙"
        "材质经设计单位优化调整。）",

        "（注：工作联系单一第6条“教学楼后侧和左侧220平方米EPDM地板取消该工程量”，"
        "该部分已并入本单第4条活动场EPDM改悬浮地板的整体调整中，不再单列。）",

        "根据现场施工情况，新建筑背面新增挡土墙18m，墙顶相对标高为0.46～1.8m；拆除旧"
        "挡土墙长65m总245m³；新建3.5m高挡土墙长40m；新建2m高挡土墙总长24m。3.5m高、"
        "2.0m高挡土墙做法分别参照图集17J008-211P-FJDB相应做法；（对应工作联系单一第7、"
        "8条挡土墙延长及拆除重建内容，具体延长米数以现场施工资料据实结算为准。）",

        "由于现场将化粪池埋置在场地西北角，致DN200污水管总长度变更为108.67m，污水检查"
        "井总数变更为8个。（对应工作联系单二第1条）",

        "原大门口整片黄锈石铺装取消1633.9㎡，下部仅做现浇硬化面积951㎡，硬化区域现调整"
        "硬化为200厚级配碎石+200厚C25混凝土（拉毛。切缝宽6mm，深5cm.填充沥青砂浆.棉毡"
        "养护）。（对应工作联系单二第2条）",

        "由于市政接入点改变，原主电缆YJV22-4X35/SC50/FC线路总长度28.65m变更为114.45m，"
        "并增设一个手孔井做法不变。（对应工作联系单二第3条）",

        "场地中央原沙池变更为戏水池，尺寸不变。（对应工作联系单二第4条）",

        "大门坡底截水沟总长度变更为28m；东侧挡土墙底部新增排水沟37m，砼制水沟盖板"
        "500*400*55厚；东侧挡土墙底新增100厚C25混凝土硬化17.48㎡。（对应工作联系单二"
        "第5条）",

        "拆除大门前停车场损坏的彩砖，并新建彩砖：规格为400*400*50厚，总面积47㎡。"
        "（对应工作联系单二第6条）",

        "根据使用单位要求新购置录像机、POE交换机各1台。（对应工作联系单二第7条）",

        "根据现场实际施工情况取消场地中央挡土墙二外立面的火山岩装饰层133.45㎡；取消"
        "沙水区的遮阳棚75㎡。（对应工作联系单二第8条）",

        "根据现场实际施工情况，将场地中间挡土墙顶的玻璃栏杆护栏做法由236.85m变更为"
        "不锈钢栏杆扶手护栏220m。做法参照东升学校场地中间坡道挡土墙两侧不锈钢护栏做法。"
        "（对应工作联系单二第9条）",
    ],
    output_path=f"{OUT_DIR}/上埇幼儿园_工程量确认单（新）.docx",
)

# ====================================================================
# 3. 东红幼儿园
# ====================================================================
build_lianxidan_2(
    school="东红幼儿园",
    use_unit="琼海市大路镇中心幼儿园东红分园",
    lianxi_bianhao="2",
    ref_bianhao="1",
    intro_text=(
        "根据现场实际情况及使用单位提出的若干意见，对原设计图纸做出部分修改，"
        "以下内容为对工作联系单（编号1）的补充说明，工作联系单（编号1）与本联系单"
        "（编号2）合并后与本项目东红幼儿园工程量确认单内容一一对应："
    ),
    items=[
        "保安亭布局调整为宽敞方案（做法同东升幼儿园保安亭）的同时，取消原设计人行"
        "通道不锈钢门；根据实际施工内容取消屋顶光伏系统。",
        "根据使用单位要求新购置录像机、POE交换机各1台。",
        "教学楼门口增加铸铁盖板水沟22.5m。内径：25cm宽，20cm高，两侧砌单砖抹灰。",
    ],
    output_path=f"{OUT_DIR}/东红幼儿园_工作联系单（二）.docx",
)

build_quedingdan_new(
    school="东红幼儿园",
    project_name_full=f"{PROJECT_FULL_NAME}-大路镇东红幼儿园室外配套附属设施建设工程",
    content_name="东红幼儿园增加部分改造工程",
    position="东红幼儿园",
    old_bianhao="002",
    note_text=(
        "说明：本确认单根据《工作联系单》（编号1，已签字盖章）与《工作联系单》"
        "（编号2，补充）整理调整，替代原《工程量确认单》计量编号002（东红幼儿园）。"
        "条目顺序已按工作联系单顺序重排，与工作联系单内容不一致或联系单中已取消的内容，"
        "均在对应条目下加注说明。"
    ),
    shiyou_text="关于东红幼儿园增加部分改造工程的确认事宜",
    items=[
        "原设计保安亭布局较紧凑，变更为布局更宽敞的方案，做法同东升幼儿园保安亭，并"
        "取消人行通道不锈钢门；根据实际施工内容取消屋顶光伏系统；（对应工作联系单一"
        "第1条，其中“取消人行通道不锈钢门”“取消屋顶光伏系统”为工作联系单二第1条"
        "补充内容）",

        "（注：工作联系单一第2条为“原设计曲面山包，变更为成品攀爬玩具”，经现场核实，"
        "最终做法为取消原设计曲面山包，面层改为铺设悬浮式地板(已含在总面积内）；以本单"
        "第4条为准，工作联系单一第2条内容不再执行。）",

        "原设计升旗台位置位于南边，变更为东边；（对应工作联系单一第3条）",

        "整个场地硬化铺装仅保留两种做法（绿化场地不变），即整体硬化+局部铺设悬浮式"
        "地板。原设计悬浮地板规格25cm×25cm×1.3cm、共721.78㎡，变更为30cm×30cm×"
        "1.6cm，铺设总面积833.16㎡；其余铺装场地均做硬化，硬化工序优化：拉毛、切缝宽"
        "6mm、深5cm、填充沥青砂浆、棉毡养护；取消原设计曲面山包，面层改为铺设悬浮式"
        "地板(已含在总面积内）；（对应工作联系单一第4条，并吸纳原第2条曲面山包变更"
        "内容）",

        "隔油池至市政路边新增φ200波纹管排水管总长39.97m，并新增3个φ400铸铁井盖检查"
        "井；保安室南侧增设1.5立方成品玻璃钢化粪池；变更主楼化粪池至市政路原DN200污水"
        "管为DN300波纹管，总长30.71m；（对应工作联系单一第5条隔油池及化粪池排放管道"
        "内容）",

        "（注：工作联系单一第6条“原设计中排水管道的尾水口位于大门围墙右外侧，周边"
        "缺乏相应的接入点，建议将管道延长至原定的排污口位置”，已通过本单第5条隔油池"
        "排水管道方案一并实施，不再单独列项。）",

        "根据使用单位要求新购置录像机、POE交换机各1台。（对应工作联系单二第2条）",

        "教学楼门口增加铸铁盖板水沟22.5m。内径：25cm宽，20cm高，两侧砌单砖抹灰。"
        "（对应工作联系单二第3条）",
    ],
    output_path=f"{OUT_DIR}/东红幼儿园_工程量确认单（新）.docx",
)

# ====================================================================
# 4. 南俸幼儿园（联系单署名为"石壁镇中心幼儿园"）
# ====================================================================
build_lianxidan_2(
    school="南俸幼儿园",
    use_unit="琼海市石壁镇中心幼儿园",
    lianxi_bianhao="2",
    ref_bianhao="1",
    intro_text=(
        "根据现场实际情况及使用单位提出的若干意见，对原设计图纸做出部分修改，"
        "以下内容为对工作联系单（编号1）的补充说明，工作联系单（编号1）与本联系单"
        "（编号2）合并后与本项目南俸幼儿园工程量确认单内容一一对应："
    ),
    items=[
        "沿着红线点1-2-3-4-5-6-7-8新建围墙224.86m的同时，取消原设计LOGO墙4.5m³。",
        "应使用单位要求取消场地东南角的攀爬区66.1㎡（与另一处攀爬活动区取消为不同"
        "位置，需分别列出）。",
        "绿化位置增加成品木栅栏。",
    ],
    output_path=f"{OUT_DIR}/南俸幼儿园_工作联系单（二）.docx",
)

build_quedingdan_new(
    school="南俸幼儿园",
    project_name_full=f"{PROJECT_FULL_NAME}-南俸幼儿园室外配套附属设施建设工程",
    content_name="南俸幼儿园增加部分改造工程",
    position="南俸幼儿园",
    old_bianhao="001",
    note_text=(
        "说明：本确认单根据《工作联系单》（编号1，已签字盖章，署名使用单位为“琼海市"
        "石壁镇中心幼儿园”，即南俸幼儿园）与《工作联系单》（编号2，补充）整理调整，"
        "替代原《工程量确认单》计量编号001（南俸幼儿园）。条目顺序已按工作联系单顺序"
        "重排，与工作联系单内容不一致或联系单中已取消的内容，均在对应条目下加注说明。"
    ),
    shiyou_text="关于南俸幼儿园增加部分改造工程的确认事宜",
    items=[
        "取消地块南侧大门外拆除新建的沥青路952.75㎡，及配套排水沟145.64m、路缘石"
        "180m.同时增设35m成品双波形防护栏；（对应工作联系单一第1条，其中防护栏材质"
        "经设计单位优化调整为双波形防护栏）",

        "原设计的沙水区域，变更为戏水池，而周边原本铺设的鹅卵石171.93㎡，亦替换为"
        "大理石地面143.5㎡；（对应工作联系单一第2条）",

        "原设计趣味花架1个，变更为沙池1个，再把原设计的探索长廊安装位置改至沙池区域；"
        "（对应工作联系单一第3条）",

        "原设计攀爬活动区66.1㎡，取消施工；（对应工作联系单一第4条）",

        "原设计的EPDM地板取消1470.61㎡，硬化区域现调整硬化为200厚级配碎石+200厚C25"
        "混凝土（工序：拉毛.切缝宽6mm,深5cm.填充沥青砂浆.棉毡养护）1697.85㎡，活动场"
        "区域改为再硬化后增铺1.6厚300×300悬浮式地板1599㎡；（注：工作联系单一第5条原"
        "为“EPDM地板调整为右侧EPDM材质、中间硬化、左侧悬浮地板三个区域混合设计”，经"
        "现场核实，最终方案为原设计EPDM地板全部取消，以本条为准。）",

        "沿着红线点1-2-3-4-5-6-7-8新建围墙224.86m；（对应工作联系单一第6条“增加铁艺"
        "围栏”，确认单未明确材质，如与铁艺围栏有出入，以现场实际施工为准）",

        "原设计仿木纹砖109.6㎡，变更为与现场相邻地板相同的做法；（对应工作联系单一"
        "第7条）",

        "另增加8m高砼电线杆1根；变压器周边增设成品隔离护栏18m；（对应工作联系单一"
        "第8条；注：联系单一第8条提到的变压器容量250KVA及高压计量设备为设备参数，"
        "不纳入本次工程量确认范围）",

        "取消LOGO墙4.5m³；（对应工作联系单二第1条）",

        "应使用单位要求取消场地东南角的攀爬区66.1㎡；（对应工作联系单二第2条）",

        "绿化位置增加成品木栅栏。（对应工作联系单二第3条）",
    ],
    output_path=f"{OUT_DIR}/南俸幼儿园_工程量确认单（新）.docx",
)

# ====================================================================
# 5. 千秋幼儿园（联系单一为使用方补充提供的截图件，署名"塔洋镇中心幼儿园"）
# ====================================================================
build_lianxidan_2(
    school="千秋幼儿园",
    use_unit="琼海市塔洋镇中心幼儿园",
    lianxi_bianhao="2",
    ref_bianhao="1",
    intro_text=(
        "根据现场实际情况及使用单位提出的若干意见，对原设计图纸做出部分修改，"
        "以下内容为对工作联系单（编号1，署名使用单位为“琼海市塔洋镇中心幼儿园”，"
        "即千秋幼儿园）的补充说明，工作联系单（编号1）与本联系单（编号2）合并后与"
        "本项目千秋幼儿园工程量确认单内容一一对应："
    ),
    items=[
        "根据现场已施工情况，取消原围墙500×500造型柱，将原2m高围墙分段成多个不同"
        "高度施工，其中2.65m高围墙18m、3.3m高围墙20m、3.05m高围墙35.22m；均为18墙"
        "双面抹灰、内外双侧一底两面真石漆涂料。",
        "按业主意见取消南向内部园门及底部台阶、两侧3m高围挡设施48㎡。",
        "根据现场已施工实际情况，取消成品木花箱4个。",
        "西侧楼梯后新增2.5m高围墙总长6m，18墙双面抹灰、内外双侧一底两面真石漆涂料。",
        "原趣味花架材料规格进行变更，均为热镀锌钢管，面喷涂咖色氟碳漆，见详图。",
        "根据使用单位要求新购置录像机、10T机械硬盘、配套显示器各1台，POE交换机1台。",
    ],
    output_path=f"{OUT_DIR}/千秋幼儿园_工作联系单（二）.docx",
)

build_quedingdan_new(
    school="千秋幼儿园",
    project_name_full=f"{PROJECT_FULL_NAME}-塔洋镇千秋幼儿园室外配套附属设施建设工程",
    content_name="千秋幼儿园增加部分改造工程",
    position="千秋幼儿园",
    old_bianhao="004",
    note_text=(
        "说明：本确认单根据《工作联系单》（编号1，署名使用单位为“琼海市塔洋镇中心"
        "幼儿园”，即千秋幼儿园）与《工作联系单》（编号2，补充）整理调整，替代原"
        "《工程量确认单》计量编号004（千秋幼儿园）。条目顺序已按工作联系单顺序重排，"
        "与工作联系单内容不一致或联系单中已取消的内容，均在对应条目下加注说明。"
    ),
    shiyou_text="关于千秋幼儿园增加部分改造工程的确认事宜",
    items=[
        "新增沙池水池顶部遮雨棚见详图；（注：工作联系单一第1条“原设计沙水池面积18"
        "平方，变更后沙池面积为15平方，水池面积为20平方”，该面积调整已按此执行）",

        "按业主意见整个场地硬化铺装仅保留两种做法(绿化场地不变）即整体硬化+局部铺设"
        "悬浮式地板，且原设计图悬浮地板规格25cm×25cm×1.3cm，655.02㎡，变更为规格为"
        "30cm×30cm×1.6cm，铺设总面积607.2㎡，其余铺装场地均简化为硬化，硬化区域现"
        "调整硬化为200厚级配碎石+200厚C25混凝土，硬化优化工序：拉毛.切缝宽6mm，深"
        "5cm.填充沥青砂浆.棉毡养护；（对应工作联系单一第2条）",

        "（注：工作联系单一第3条“原设计没有学校全称的广告牌，建议增加”，经研究"
        "未采纳，此项取消。）",

        "按业主意见取消南向内部园门及底部台阶、两侧3m高围挡设施48㎡；（对应工作联系"
        "单二第2条）",

        "根据现场已施工实际情况，取消成品木花箱4个；（对应工作联系单二第3条）",

        "根据现场已施工情况，取消原围墙500×500造型柱，将原2m高围墙分段成多个不同"
        "高度施工，其中2.65m高围墙18m、3.3m高围墙20m、3.05m高围墙35.22m；均为18墙"
        "双面抹灰、内外双侧一底两面真石漆涂料；（对应工作联系单二第1条）",

        "西侧楼梯后新增2.5m高围墙总长6m，18墙双面抹灰、内外双侧一底两面真石漆涂料；"
        "（对应工作联系单二第4条）",

        "原趣味花架材料规格进行变更，均为热镀锌钢管，面喷涂咖色氟碳漆，见详图；"
        "（对应工作联系单二第5条）",

        "根据使用单位要求新购置录像机、10T机械硬盘、配套显示器各1台，POE交换机1台。"
        "（对应工作联系单二第6条）",
    ],
    output_path=f"{OUT_DIR}/千秋幼儿园_工程量确认单（新）.docx",
)

print("全部5所学校文件生成完成")
