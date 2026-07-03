# -*- coding: utf-8 -*-
"""
生成东升幼儿园：
1. 工作联系单（二） —— 版式1:1复刻《工作联系单》原件模板
2. 工程量确认单（新）—— 版式1:1复刻《工程量确认单》原件模板（坐标已从原PDF精确提取）

仅涉及东升幼儿园（万泉镇中心幼儿园东升分园）部分。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "仿宋_GB2312"   # 原件字体为 FangSong；用户在 Word/WPS 中打开会正常显示仿宋
BODY_SIZE = 12
TITLE_SIZE = 24

CONTENT_WIDTH_CM = 17.33  # 与原件表格宽度一致（491.35pt）


def set_east_asia_font(run, size=BODY_SIZE, bold=False, name=FONT_NAME):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def set_single_spacing(paragraph, space_before=0, space_after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{tag}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        mar.append(node)
    tcPr.append(mar)


def shrink_table_borders(table):
    """确保表格边框为黑色细实线（Table Grid 默认即可，这里保留接口）"""
    table.style = "Table Grid"


def new_section_margins(doc, top=1.6, bottom=1.6, left=1.84, right=1.84):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_title(doc, text, space_after_pt=22.65 * 0.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p, space_before=0, space_after=space_after_pt)
    run = p.add_run(text)
    set_east_asia_font(run, size=TITLE_SIZE, bold=False)
    return p


def add_plain_para(doc, text, size=BODY_SIZE, indent_chars=0, align=None,
                    space_before=0, space_after=0):
    p = doc.add_paragraph()
    set_single_spacing(p, space_before, space_after)
    if indent_chars:
        p.paragraph_format.first_line_indent = Pt(size * indent_chars)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_east_asia_font(run, size=size)
    return p


def set_cell(cell, text, size=BODY_SIZE, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             valign=WD_ALIGN_VERTICAL.CENTER, multi_line=None):
    cell.text = ""
    set_cell_margins(cell)
    cell.vertical_alignment = valign
    lines = multi_line if multi_line else [text]
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        set_single_spacing(p, 0, 0)
        p.alignment = align
        run = p.add_run(line)
        set_east_asia_font(run, size=size, bold=bold)
    return cell


def set_row_height(row, cm_value, rule=WD_ROW_HEIGHT_RULE.AT_LEAST):
    row.height = Cm(cm_value)
    row.height_rule = rule


# ------------------------------------------------------------------
# 1. 工作联系单（二）—— 复刻《工作联系单》原件版式
# ------------------------------------------------------------------

def build_lianxidan_2():
    doc = Document()
    new_section_margins(doc)

    add_title(doc, "工作联系单", space_after_pt=14)

    # 工程名称 + 编号（表格外的单独一行，编号靠右）
    p = doc.add_paragraph()
    set_single_spacing(p, 0, 10)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(CONTENT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(
        "工程名称：东升、东红、千秋、东平、南俸、上埇和机关等7所幼儿园室外配套附属设施建设工程"
        "\t编号：2"
    )
    set_east_asia_font(run, size=BODY_SIZE)

    # 外框大表格：第1行合并为致/事由/条款/签字区，第2~6行为五个意见栏
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col0_w, col1_w = Cm(6.0), Cm(CONTENT_WIDTH_CM - 6.0)
    for row in table.rows:
        row.cells[0].width = col0_w
        row.cells[1].width = col1_w

    # ---- 第1行：合并单元格，放致/事由/正文/签字 ----
    top_cell = table.cell(0, 0).merge(table.cell(0, 1))
    top_cell.text = ""
    set_cell_margins(top_cell, top=120, bottom=120, left=140, right=140)
    top_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    def top_para(first=False):
        p = top_cell.paragraphs[0] if first else top_cell.add_paragraph()
        set_single_spacing(p, 0, 0)
        return p

    # 致：单位列表（悬挂缩进 + Tab 对齐右侧单位类型标签）
    p = top_para(first=True)
    pf = p.paragraph_format
    pf.left_indent = Cm(1.1)
    pf.first_line_indent = Cm(-1.1)
    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Cm(7.2))
    units = [
        ("致：海口市工程监理有限公司", "（监理单位）"),
        ("琼海市万泉镇中心幼儿园", "（使用单位）"),
        ("中昌设计集团有限公司", "（设计单位）"),
        ("琼海市城市投资运营有限公司", "（建设单位）"),
        ("琼海市教育局", ""),
    ]
    for i, (name, label) in enumerate(units):
        run = p.add_run(f"{name}\t{label}")
        set_east_asia_font(run, size=BODY_SIZE)
        if i != len(units) - 1:
            run.add_break()

    p = top_para()
    set_single_spacing(p, 8, 0)
    run = p.add_run("事由：")
    set_east_asia_font(run, size=BODY_SIZE)

    p = top_para()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(BODY_SIZE * 2)
    run = p.add_run(
        "根据现场实际情况及使用单位提出的若干意见，对原设计图纸做出部分修改，"
        "以下内容为对工作联系单（编号1）的补充说明，工作联系单（编号1）与本联系单"
        "（编号2）合并后与本项目东升幼儿园工程量确认单内容一一对应："
    )
    set_east_asia_font(run, size=BODY_SIZE)

    items = [
        "1.大门口消防车道拆除原混凝土路面后新建，面积103.79㎡（拆除量按施工资料结算）；"
        "大门内消防车道坡道两侧增加挡土墙，总长22.97m；坡道两侧新增镀锌栏杆扶手，总长度15.5m。",

        "2.取消原电动伸缩门1座，替换为成品双扇咖色氟碳漆热镀锌方管大门，高2.5m总宽6m；"
        "大门外指定场地换填300厚种植土铺设台湾草438㎡，种植面积以最终现场为准。",

        "3.应使用单位要求拆除单层混合结构建筑(面积据拆除影像资料据实结算）；清理后原场地做"
        "硬化并铺设悬浮式地板，面积70平方米，拼接图案厂家提供，中心点标高高于周边地面0.03"
        "并与之找坡平接；原单层混合结构建筑门口位置增设铸铁盖板14.4m。",

        "4.应业主要求对新建围墙做法进行变更，其中总高2.9m的围墙长105.05m，"
        "总高2.5m的围墙长29.94m，做法见图。",

        "5.根据现场排水需要，在隔油池（见工作联系单一第3条）旁边增设两个铸铁篦子雨水口，"
        "通过5.4mDN200波纹管排至雨水井。",
    ]
    for text in items:
        p = top_para()
        run = p.add_run(text)
        set_east_asia_font(run, size=BODY_SIZE)

    p = top_para()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(BODY_SIZE * 2)
    set_single_spacing(p, 6, 0)
    run = p.add_run("请各参建单位予以确认。")
    set_east_asia_font(run, size=BODY_SIZE)

    for label in ("承包单位：", "项目经理：", "日\u3000\u3000期："):
        p = top_para()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_single_spacing(p, 4, 0)
        pf = p.paragraph_format
        pf.right_indent = Cm(2.0)
        run = p.add_run(label)
        set_east_asia_font(run, size=BODY_SIZE)

    # ---- 第2~6行：五个意见栏（左：单位意见；右：盖章/签字/日期） ----
    opinion_rows = [
        ("监理单位意见：", ["监理单位（盖章）：", "总监理工程师（签字）：", "日期："]),
        ("使用单位意见：", ["使用单位（盖章）：", "项目负责人（签字）：", "日期："]),
        ("设计单位意见：", ["设计单位（盖章）：", "项目负责人（签字）：", "日期："]),
        ("建设单位意见：", ["建设单位（盖章）：", "项目负责人（签字）：", "日期："]),
        ("教育局意见：", ["琼海市教育局（盖章）：", "项目负责人（签字）：", "日期："]),
    ]
    for i, (left_label, right_lines) in enumerate(opinion_rows, start=1):
        row = table.rows[i]
        set_row_height(row, 2.7)
        set_cell(row.cells[0], "", valign=WD_ALIGN_VERTICAL.TOP, multi_line=[left_label])
        set_cell(row.cells[1], "", valign=WD_ALIGN_VERTICAL.TOP, multi_line=right_lines)

    doc.save("/workspace/output_docs/东升幼儿园_工作联系单（二）.docx")
    print("联系单二 done")


# ------------------------------------------------------------------
# 2. 工程量确认单（新）—— 坐标已从原PDF精确提取，1:1还原表格结构
# ------------------------------------------------------------------

def build_quedingdan_new():
    doc = Document()
    new_section_margins(doc)

    add_title(doc, "工程量确认单", space_after_pt=16)

    # 说明文字（不属于原模板，放在标题下、表格上，字号略小以区分为整理说明）
    p = doc.add_paragraph()
    set_single_spacing(p, 0, 10)
    run = p.add_run(
        "说明：本确认单根据《工作联系单》（编号1，已签字盖章）与《工作联系单》"
        "（编号2，补充）整理调整，替代原《工程量确认单》计量编号003（东升幼儿园）。"
        "条目顺序已按工作联系单顺序重排，与工作联系单内容不一致或联系单中已取消的内容，"
        "均在对应条目下加注说明。"
    )
    set_east_asia_font(run, size=10.5)

    # ---- 顶部信息表：2行 x 4列，比例还原自原PDF坐标 ----
    info = doc.add_table(rows=2, cols=4)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    widths = [Cm(2.88), Cm(7.31), Cm(1.67), Cm(5.47)]
    for row in info.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    set_row_height(info.rows[0], 3.27)
    set_row_height(info.rows[1], 1.69)

    set_cell(info.rows[0].cells[0], "工程名称", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(
        info.rows[0].cells[1], "",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        valign=WD_ALIGN_VERTICAL.TOP,
        multi_line=[
            "东升、东红、千秋、东平、南俸、",
            "上埇和机关等7所幼儿园室外配套",
            "附属设施建设工程-万泉镇中心幼",
            "儿园东升分园室外配套附属设施建",
            "设工程",
        ],
    )
    set_cell(info.rows[0].cells[2], "内容", align=WD_ALIGN_PARAGRAPH.CENTER,
              valign=WD_ALIGN_VERTICAL.TOP)
    set_cell(info.rows[0].cells[3], "", align=WD_ALIGN_PARAGRAPH.LEFT,
              valign=WD_ALIGN_VERTICAL.TOP,
              multi_line=["东升幼儿园增加部分改", "造工程"])

    set_cell(info.rows[1].cells[0], "位置", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(info.rows[1].cells[1], "东升幼儿园", align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(info.rows[1].cells[2], "", align=WD_ALIGN_PARAGRAPH.CENTER,
              multi_line=["计量", "编号"])
    set_cell(info.rows[1].cells[3], "003（修订）", align=WD_ALIGN_PARAGRAPH.LEFT)

    # ---- 事由 + 正文（紧贴表格下方，几乎无间距，与原件一致） ----
    add_plain_para(doc, "事由：关于东升幼儿园增加部分改造工程的确认事宜",
                   space_before=6, space_after=0)
    add_plain_para(doc, "根据会议纪要，结合使用单位提出的若干意见做出修改，修改内容如下：",
                   indent_chars=2, space_before=0, space_after=4)

    items = [
        "1.原设计PVC运动地板735.22㎡,变更为减震悬浮地板规格和型号：30cm×30cm×1.6cm，"
        "铺设面积664㎡，硬化区域现调整硬化为200厚级配碎石+200厚C25混凝土725㎡，硬化优"
        "化工序：拉毛.切缝宽6mm,深5cm.填充沥青砂浆.棉毡养护；（对应工作联系单一第1条）",

        "2.原设计靠近新建教学楼位置的砖砌围墙，变更为镀锌管栏杆围墙(小柱："
        "50mm×50mm×2.0mm@80; 大柱：80mm×80mm×2.5mm@3000；横杆上下两道规格"
        "40×40×2.5mm）, 立柱整体高度2m，长度124米;"
        "（注：工作联系单一第2条原为“铝艺栏杆围墙，立柱整体高度2.5米”，"
        "经设计单位优化调整，现变更为本条内容。）",

        "3.增设一个容量为2T的成品玻璃钢隔油池，埋置16.79mDN200波纹管接至W8污水井；"
        "（注：工作联系单一第3条中“沙池旁增加两个洗手洗脚水龙头并进行硬化处理”"
        "“戏水池旁鹅卵石小径改为大理石小径”“教学楼入口处花池内增设排水沟”"
        "“大门绿化位置优化缩小”等内容，现均已取消。）",

        "4.（注：工作联系单一第4条“新建戏水池顶部加装顶棚（防止树枝、树叶、"
        "昆虫掉入水中）”，此项已取消。）",

        "5.旁边增设两个铸铁篦子雨水口，通过5.4mDN200波纹管排至雨水井；"
        "（对应工作联系单二第5条，为第3条隔油池排水配套新增内容）",

        "6.大门口消防车道，面积103.79㎡(拆除原混凝土路面后新建，拆除量按施工资料）；大门"
        "内消防车道坡道两侧增加挡土墙总长22.97m，坡道两侧新增镀锌栏杆扶手，总长度15.5m；"
        "（对应工作联系单二第1条）",

        "7.取消原电动伸缩门1座，替换为成品双扇咖色氟碳漆热镀锌方管大门，高2.5m总宽6m；大"
        "门外指定场地换填300厚种植土铺设台湾草438㎡，种植面积以最终现场为准；"
        "（对应工作联系单二第2条）",

        "8.应使用单位要求拆除单层混合结构建筑(面积据拆除影像资料据实结算）；清理后原场地做"
        "硬化并铺设悬浮式地板，面积70平方米，拼接图案厂家提供，中心点标高高于周边地面0.03"
        "并与之找坡平接；原单层混合结构建筑门口位置铸铁盖板14.4m。"
        "（对应工作联系单二第3条）",

        "9.应业主要求对新建围墙做法进行变更，其中总高2.9m的围墙长105.05m，总高2.5m的围墙"
        "长29.94m，做法见图。（对应工作联系单二第4条）",
    ]
    for text in items:
        add_plain_para(doc, text, space_before=0, space_after=4)

    add_plain_para(doc, "以上变更工程量详见后附计算表。", space_before=6, space_after=8)

    # ---- 底部四方意见栏：单行、四列，标签左上对齐，其余留空 ----
    bottom = doc.add_table(rows=1, cols=4)
    bottom.style = "Table Grid"
    bottom.alignment = WD_TABLE_ALIGNMENT.CENTER
    bottom.autofit = False
    widths2 = [Cm(4.19), Cm(4.20), Cm(4.47), Cm(4.48)]
    labels = ["建设单位意见：", "代管单位意见：", "监理单位意见：", "施工单位意见："]
    for cell, w in zip(bottom.rows[0].cells, widths2):
        cell.width = w
    set_row_height(bottom.rows[0], 4.32)
    for cell, label in zip(bottom.rows[0].cells, labels):
        set_cell(cell, label, valign=WD_ALIGN_VERTICAL.TOP)

    doc.save("/workspace/output_docs/东升幼儿园_工程量确认单（新）.docx")
    print("新确认单 done")


if __name__ == "__main__":
    build_lianxidan_2()
    build_quedingdan_new()
