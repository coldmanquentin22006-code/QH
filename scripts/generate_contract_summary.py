# -*- coding: utf-8 -*-
"""
生成《施工合同基本情况说明》Word 文档（现代简约风格）。
数据来源：2024-192.东升、东红、千秋、东平、南俸、上埇和机关等7所幼儿园
室外配套附属设施建设工程施工合同.doc
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 配色（暖色调，类似 Claude/Anthropic 品牌色系） ----------
ACCENT = RGBColor(0xC1, 0x5A, 0x33)       # 暖赤陶色，用于强调
ACCENT_DARK = RGBColor(0x8A, 0x3E, 0x22)  # 深赤陶色，用于价格数字
INK = RGBColor(0x2B, 0x28, 0x24)          # 近黑正文色
SUBTLE = RGBColor(0x8A, 0x82, 0x78)       # 灰褐色，用于辅助文字
LINE_GRAY = "E3DDD4"                       # 分隔线/边框（十六进制，无#）
CARD_BG = "FBF3EC"                         # 信息卡浅底色
HEADER_BG = "F2E4D8"                       # 表头浅底色
ACCENT_HEX = "C15A33"
ACCENT_DARK_HEX = "8A3E22"

BODY_FONT = "微软雅黑"
HEADING_FONT = "微软雅黑"


def set_font(run, font_name=BODY_FONT, size=Pt(11), bold=False, color=INK, italic=False):
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side, spec in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        if spec is None:
            continue
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), spec.get("val", "single"))
        el.set(qn("w:sz"), str(spec.get("sz", 4)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), spec.get("color", LINE_GRAY))
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins(cell, top=80, bottom=80, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def add_space(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("")
    r.font.size = Pt(pt * 0.6)


def add_title_block(doc, eyebrow, title, meta_lines):
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = Pt(4)
    r0 = p0.add_run(eyebrow)
    set_font(r0, BODY_FONT, Pt(10.5), bold=True, color=ACCENT)
    r0.font.italic = False

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(10)
    r1 = p1.add_run(title)
    set_font(r1, HEADING_FONT, Pt(23), bold=True, color=INK)

    # 分隔线（底部细边框）
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_after = Pt(10)
    pPr = rule_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(2)
    for i, line in enumerate(meta_lines):
        if i > 0:
            sep = meta_p.add_run("    |    ")
            set_font(sep, BODY_FONT, Pt(9.5), color=SUBTLE)
        r = meta_p.add_run(line)
        set_font(r, BODY_FONT, Pt(9.5), color=SUBTLE)


def add_section_heading(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True

    num_run = p.add_run(f"{number}  ")
    set_font(num_run, HEADING_FONT, Pt(13.5), bold=True, color=ACCENT)
    text_run = p.add_run(text)
    set_font(text_run, HEADING_FONT, Pt(13.5), bold=True, color=INK)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LINE_GRAY)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_body_text(doc, text, size=Pt(10.5), color=INK, space_after=8, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_font(run, BODY_FONT, size, color=color)
    return p


def build_info_table(doc, rows, key_width_cm=3.6):
    """构建无边框、斑马纹、卡片风格的信息表。"""
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(table)
    table.autofit = False

    thin_bottom = {"val": "single", "sz": 4, "color": LINE_GRAY}

    for idx, (key, value) in enumerate(rows):
        row_cells = table.add_row().cells
        row_cells[0].width = Cm(key_width_cm)
        row_cells[1].width = Cm(16.5 - key_width_cm)

        bg = CARD_BG if idx % 2 == 0 else "FFFFFF"
        for c in row_cells:
            shade_cell(c, bg)
            set_cell_margins(c, top=100, bottom=100, left=180, right=180)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(c, bottom=thin_bottom)

        k_p = row_cells[0].paragraphs[0]
        k_p.paragraph_format.line_spacing = 1.3
        k_run = k_p.add_run(key)
        set_font(k_run, BODY_FONT, Pt(10), bold=True, color=SUBTLE)

        v_p = row_cells[1].paragraphs[0]
        v_p.paragraph_format.line_spacing = 1.3
        value_lines = value.split("\n")
        v_run = v_p.add_run(value_lines[0])
        set_font(v_run, BODY_FONT, Pt(10.5), bold=False, color=INK)
        for extra_line in value_lines[1:]:
            v_run.add_break()
            v_run2 = v_p.add_run(extra_line)
            set_font(v_run2, BODY_FONT, Pt(10.5), bold=False, color=INK)

    return table


def add_price_highlight(doc, label, amount, sub_items):
    """突出显示合同总价的强调卡片。"""
    table = doc.add_table(rows=1, cols=1)
    no_table_borders(table)
    cell = table.rows[0].cells[0]
    shade_cell(cell, ACCENT_HEX)
    set_cell_margins(cell, top=260, bottom=260, left=320, right=320)

    p_label = cell.paragraphs[0]
    p_label.paragraph_format.space_after = Pt(4)
    r_label = p_label.add_run(label)
    set_font(r_label, BODY_FONT, Pt(10.5), bold=True, color=RGBColor(0xFF, 0xF3, 0xEA))

    p_amount = cell.add_paragraph()
    p_amount.paragraph_format.space_after = Pt(6)
    r_amount = p_amount.add_run(amount)
    set_font(r_amount, HEADING_FONT, Pt(26), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for sub in sub_items:
        p_sub = cell.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(1)
        r_sub = p_sub.add_run(f"·  {sub}")
        set_font(r_sub, BODY_FONT, Pt(9.5), color=RGBColor(0xFC, 0xE4, 0xD6))

    return table


def add_note_item(doc, index, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.left_indent = Cm(0.4)
    num_run = p.add_run(f"{index}   ")
    set_font(num_run, BODY_FONT, Pt(9.5), bold=True, color=ACCENT)
    text_run = p.add_run(text)
    set_font(text_run, BODY_FONT, Pt(9.5), color=SUBTLE)


# ================= 生成文档 =================
doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(1.9)
section.bottom_margin = Cm(1.9)

style = doc.styles["Normal"]
style.font.name = BODY_FONT
style.font.size = Pt(10.5)
style_rpr = style.element.get_or_add_rPr()
style_rfonts = style_rpr.find(qn("w:rFonts"))
if style_rfonts is None:
    style_rfonts = style_rpr.makeelement(qn("w:rFonts"), {})
    style_rpr.append(style_rfonts)
style_rfonts.set(qn("w:eastAsia"), BODY_FONT)

# ---- 标题区 ----
add_title_block(
    doc,
    eyebrow="工程合同 · 情况说明",
    title="7所幼儿园室外配套附属设施建设工程",
    meta_lines=["合同编号 QHCT(FW)2024-192号", "2026年7月3日"],
)

add_body_text(
    doc,
    "为便于领导快速掌握该项目施工合同的核心信息，现将合同名称、签约双方、"
    "合同价格等要点整理如下，详细条款请见合同全文。",
    size=Pt(10.5),
    color=SUBTLE,
    space_after=14,
)

# ---- 一、项目基本情况 ----
add_section_heading(doc, "01", "项目基本情况")
build_info_table(
    doc,
    [
        ("合同编号", "QHCT(FW)2024-192号"),
        ("项目名称", "东升、东红、千秋、东平、南俸、上埇和机关等7所幼儿园\n室外配套附属设施建设工程"),
        ("工程地点", "琼海市嘉积镇、万泉镇、塔洋镇、大路镇、会山镇及石壁镇"),
        ("资金来源", "市级财政资金"),
        ("计划工期", "180个日历天"),
        ("签订地点", "海南省琼海市"),
    ],
)
add_space(doc, 8)

# ---- 二、合同双方 ----
add_section_heading(doc, "02", "合同双方")
build_info_table(
    doc,
    [
        ("发包人（业主）", "琼海市城市投资运营有限公司"),
        ("发包人地址", "琼海市嘉积镇爱华东路19号"),
        ("发包人信用代码", "91469002747790848Q"),
        ("承包人（施工单位）", "陕西宏基源建设有限公司"),
        ("承包人地址", "西安市文艺北路5号敬业大厦1302室—1307室"),
        ("承包人信用代码", "91610000741282199A"),
        ("项目经理", "樊俊飞（注册证号：陕1612022202302782）"),
    ],
)
add_space(doc, 8)

# ---- 三、合同价格 ----
add_section_heading(doc, "03", "合同价格")
add_price_highlight(
    doc,
    label="签约合同价（含增值税）",
    amount="¥ 11,444,630.26",
    sub_items=[
        "不含税价：¥ 10,499,660.79",
        "增值税额（税率9%）：¥ 944,969.47",
    ],
)
add_space(doc, 8)
build_info_table(
    doc,
    [
        ("合同价格形式", "固定单价合同"),
        ("下浮率", "4.27%"),
        ("安全文明施工费", "¥ 429,538.91"),
        ("专业工程暂估价", "¥ 200,537.43"),
    ],
)
add_space(doc, 8)

# ---- 四、备注 ----
add_section_heading(doc, "04", "备注")
add_note_item(
    doc,
    "1",
    "以上信息摘录自《建设工程施工合同》（合同编号：QHCT(FW)2024-192号）"
    "第一部分“合同协议书”，最终结算合同价以工程竣工验收结算审核结果为准。",
)
add_note_item(doc, "2", "如需查阅合同全文或相关附件（招标文件、投标文件、图纸等），请与项目经办人联系。")

# ---- 落款 ----
add_space(doc, 20)
tail = doc.add_paragraph()
tail.alignment = WD_ALIGN_PARAGRAPH.RIGHT
tail.paragraph_format.space_after = Pt(2)
tail_run = tail.add_run("汇报人：____________")
set_font(tail_run, BODY_FONT, Pt(10.5), color=SUBTLE)

tail2 = doc.add_paragraph()
tail2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
tail2_run = tail2.add_run("日期：2026年7月3日")
set_font(tail2_run, BODY_FONT, Pt(10.5), color=SUBTLE)

output_path = "合同基本情况说明.docx"
doc.save(output_path)
print(f"已生成：{output_path}")
